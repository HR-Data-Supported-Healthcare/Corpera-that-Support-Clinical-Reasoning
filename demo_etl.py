"""
Basic ETL proces to create a corpora ready for annotation (NER) step.
Unoptimized.
docx data should be stored in `DATA_LOC` destination for file to work
"""

from docx import Document
import os
from datetime import datetime
import jsonpickle

#CONSTANTS
#TODO Implement flags
LEMMATIZATION_FLAG = False
STEMMING_FLAG = False
STOP_WORDS_FLAG = False
DATA_LOC =  "./data"

#PUBLIC VARIABLES

#FUNCTIONS
def Extract() -> list[Document]:
    doc_list : list [Document] = []
    for filename in os.listdir(DATA_LOC):
        doc_list.append((Document(f"{DATA_LOC}/{filename}"), filename))
    return doc_list

def Transform(doc_list):
        corpora_obj = {}
        for current_doc, filename in doc_list:
            print(len(current_doc.paragraphs))
            corpora : list[list] = []
            i: int = 0
            while i < len(current_doc.paragraphs):
                #Remove Titlepage
                if current_doc.paragraphs[i].style.name.startswith("Title"):
                     print("Removing paragraph: Title")
                     pass
        
                #Remove table of contents heading + body
                elif current_doc.paragraphs[i].text.lower().startswith("inhoudsopgave") or current_doc.paragraphs[i].style.name == "TOCHeading":
                    print("Removing paragraph: toc")
                    #TODO: Perhaps table of contents is merged in one paragraph, should check
                    #Remove next paragraph too
                    i += 1

                #Remove paragraph headings
                elif current_doc.paragraphs[i].style.name.startswith("Heading"):
                     print("Removing paragraph: heading")
                     #TODO: Perhaps headings can be transformed instead -> may hold value
                     pass
                
                #Remove source list 
                elif "bibliography" in current_doc.paragraphs[i].style.name.lower() or current_doc.paragraphs[i].text.lower().startswith("literatuurlijst"):
                     print("Removing paragraph: source list")
                     pass
                
                #Remove table annotations
                elif False:
                     print("Removing paragraph: annotation")
                     #TODO
                     pass
                
                #Remove picture annotations
                elif False:
                     print("Removing paragraph: annotation")
                     #TODO
                     pass
                
                #Remove empty paragraphs
                elif current_doc.paragraphs[i].text == "":
                     print("Removing paragraph: empty paragraph") 
                     #TODO
                     pass
                
                else:
                    print("Adding text to corpora")
                    corpora.append([current_doc.paragraphs[i].text, []])
                i += 1
            corpora_obj[filename] = corpora
        return corpora_obj

def Load(dest_str: str, data):
    json_str = jsonpickle.encode(data)
    with open (dest_str, "w") as f:
         f.write(json_str)

if __name__ == "__main__":
    extracted_data = Extract()
    transformed_data = Transform(extracted_data)
    load_dir = "./output"
    Load(f"{load_dir}/Demo {datetime.now().strftime('%m-%d-%Y,%H-%M-%S')}.json", transformed_data)