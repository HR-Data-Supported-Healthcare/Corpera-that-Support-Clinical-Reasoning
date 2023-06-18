from docx import Document
from docx.shared import RGBColor
from modifiers.paragraph_identifier import ParagraphIdentifier
from modifiers.text_modifier import TextModifier
from data_classes.flag import VisualizeFlags
from data_classes.paragraph import Paragraph
import json

@staticmethod
def visualize_text_modifications(paragraph_list: list[Paragraph], vis_flags: VisualizeFlags, stop_words_list: list[str] = []) -> Document:
    """
    Exports `docx` document containing text paragraphs, coloring paragraphs based on `TextModifier` criteria and given flags.
    Also colors stop words

    Args:
        paragraph_list (list[Paragraph]): List of paragraphs to apply modifications on and generate docx document from
        vis_flags (VisualizeFlags): Configuration flags for paragraph and text modifications (which ones should be colored and applied).
        stop_words_list (list[str]) (optional, default = []): stop words to color in paragraphs 
    if vis_flags.lemmatization and vis_flags.stemming both true, only does lemmatization
    """
    #colors for each paragraph to differentiate
    stop_words_color = RGBColor(255, 255, 255) #White
    headings_color = RGBColor(0, 0, 255) #Blue
    titlepage_color = RGBColor(179, 46, 161) #Purple
    bibliography_color = RGBColor(213, 165, 68) #Brown
    table_figures_color = RGBColor(102, 0, 51) #Dark purple
    visualized_document = Document()

    if vis_flags.visualize_lemmatization:
        stop_words_list = [TextModifier.apply_lemmatization(w) for w in stop_words_list]
        
    elif vis_flags.visualize_stemming:
        stop_words_list = [TextModifier.apply_stemming(w) for w in stop_words_list]
    #remove duplicate entries in stop word list
    stop_words_list = list(set(stop_words_list))
    for paragraph in paragraph_list: 
        paragraph_text = paragraph.text
        main_color = RGBColor(0, 0, 0) #If paragraph does not conform to any removal specifications, make it black
        if ParagraphIdentifier.is_paragraph_heading(paragraph.style.name) and vis_flags.visualize_headings:
            main_color = headings_color
        elif ParagraphIdentifier.is_paragraph_title(paragraph.style.name, paragraph.text) and vis_flags.visualize_titlepage:
            main_color = titlepage_color
        elif ParagraphIdentifier.is_paragraph_table_or_figure(paragraph.text) and vis_flags.visualize_table_figures:
            main_color = table_figures_color
        elif ParagraphIdentifier.is_paragraph_bibliography(paragraph.style.name, paragraph.text) and vis_flags.visualize_bibliography:
            main_color = bibliography_color

        if vis_flags.visualize_lemmatization:
            paragraph_text = TextModifier.apply_lemmatization(paragraph_text)
        elif vis_flags.visualize_stemming:
            paragraph_text = TextModifier.apply_stemming(paragraph_text)
        new_p = visualized_document.add_paragraph()
        if len(stop_words_list) > 0:
            color_paragraph(paragraph_text , new_p, main_color, stop_words_list, stop_words_color)
        else:
            color_paragraph(paragraph_text , new_p, main_color)

    return visualized_document

def color_paragraph(text: str, paragraph, main_color: RGBColor, stop_words_list: list[str]= None, stop_words_color: RGBColor = None):
    if stop_words_list is None:
        run = paragraph.add_run(text)
        run.font.color.rgb = main_color
        return paragraph
    
    split_text = text.split()
    current_run_text : str = ""
    for word in split_text:
        if word.lower() not in stop_words_list: 
            current_run_text = current_run_text + word + " "
        else:
            if len(current_run_text) > 0:
                main_run = paragraph.add_run(current_run_text)
                main_run.font.color.rgb = main_color
                current_run_text = ""
            stop_words_run = paragraph.add_run(word + " ")
            stop_words_run.font.color.rgb = stop_words_color
    #Add the remaining words to paragraph
    remaining_run = paragraph.add_run(current_run_text)
    remaining_run.font.color.rgb = main_color
    return paragraph

#Constants
DATA_DIR = "./data/BL_casereport1.docx"
STOP_WORDS_DIR = "./stopwords.json"
OUTPUT_DIR = "./"
OUTPUT_NAME = "example.docx"

if __name__ == "__main__":
    with open (STOP_WORDS_DIR, "r") as f:
        sw = json.load(f)
    flags = VisualizeFlags(True, False, True, True, True, True, True)
    doc = Document(DATA_DIR)
    new_doc = visualize_text_modifications(doc.paragraphs, flags, sw )
    new_doc.save(f"{OUTPUT_DIR}/{OUTPUT_NAME}")