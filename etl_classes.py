from docx import Document
import jsonpickle
from docxcompose.composer import Composer
from paragraph_modifier import ParagraphModifier
from text_modifier import TextModifier
from base_etl_class import BaseETL
from json import loads

#TODO: Perhaps classes can be stored in seperate files for less total imports
#TODO: Revise class structure idea, might not be necessary

class DemoETL(BaseETL):
    """ TODO: Docstring """
    def __init__(self, data_dir: str, dest_dir: str):
        super().__init__(data_dir, dest_dir)

    def _transform(self, data):
        """ TODO: Docstring """
        corpora_obj = {}
        for current_doc, filename in data:
            print(len(current_doc.paragraphs))
            corpora : list[list] = []
            i: int = 0
            while i < len(current_doc.paragraphs):
                #Remove Titlepage
                if current_doc.paragraphs[i].style.name.startswith("Title"):
                     print("Removing paragraph: Title")
                     pass
        
                #Remove table of contents heading + body
                elif current_doc.paragraphs[i].text.lower().startswith("inhoudsopgave") or current_doc.paragraphs[i].style.name == "TOCHeading":
                    print("Removing paragraph: toc")
                    #TODO: Perhaps table of contents is merged in one paragraph, should check
                    #Remove next paragraph too
                    i += 1

                #Remove paragraph headings
                elif current_doc.paragraphs[i].style.name.startswith("Heading"):
                     print("Removing paragraph: heading")
                     #TODO: Perhaps headings can be transformed instead -> may hold value
                     pass
                
                #Remove source list 
                elif "bibliography" in current_doc.paragraphs[i].style.name.lower() or current_doc.paragraphs[i].text.lower().startswith("literatuurlijst"):
                     print("Removing paragraph: source list")
                     pass
                
                #Remove table annotations
                elif False:
                     print("Removing paragraph: annotation")
                     #TODO
                     pass
                
                #Remove picture annotations
                elif False:
                     print("Removing paragraph: annotation")
                     #TODO
                     pass
                
                #Remove empty paragraphs
                elif current_doc.paragraphs[i].text == "":
                     print("Removing paragraph: empty paragraph") 
                     #TODO
                     pass
                
                else:
                    print("Adding text to corpora")
                    corpora.append([current_doc.paragraphs[i].text, []])
                i += 1
            corpora_obj[filename] = corpora
        return corpora_obj
    
    def _load(self, dest_str: str, data):
        """
        TODO: Docstring
        """
        json_str = jsonpickle.encode(data)
        with open (dest_str, "w") as f:
            f.write(json_str)

class DynamicETL(BaseETL):
    """ TODO: Docstring """

    def __init__(self, data_dir: str, dest_dir: str, flag_dict: dict):
        self._execute_etl(data_dir, dest_dir, flag_dict)

    def _execute_etl(self, data_dir: str, dest_dir: str, flag_dict: dict) -> None:
        extracted_data = self._extract(data_dir)
        transformed_data = self._transform(extracted_data, flag_dict)
        self._load(dest_dir, transformed_data)

    def _transform(self, data: list[Document], flag_dict: dict):
        """ TODO: Docstring """
        corpora_total = {}
        corpora_total["config"] = flag_dict
        removed_paragraphs = []
        for current_doc, filename in data:
            current_doc_corpora : list[dict] = []
            doc_paragraphs = current_doc.paragraphs
            #headings should be removed separately (different paragraphs to remove together)
            if flag_dict["paragraph_filters"]["headings"]:
                corpora_with_headings_removed, removed_headings = ParagraphModifier.remove_headings(doc_paragraphs)   
                doc_paragraphs = corpora_with_headings_removed
                removed_paragraphs.extend(removed_headings)
            current_doc_corpora.extend(self.__transform_paragraphs_from_doc(doc_paragraphs, flag_dict))
            corpora_total[filename] = current_doc_corpora
        return corpora_total
            
    def __transform_paragraphs_from_doc(self, doc_paragraphs, flag_dict: dict) -> list[dict]:
        """TODO: Docstring"""
        total_doc_corpora = []
        for paragraph in doc_paragraphs:
            #For each paragraph, check the flag dictionary
            total_doc_corpora.append({})
            total_doc_corpora[-1]["text"] = paragraph.text
            total_doc_corpora[-1]["style"] = paragraph.style.name
            
            # Remove empty paragraphs
            if paragraph.text == "" or paragraph.text == "\n":
                del total_doc_corpora[-1]
                continue

            if flag_dict["text_transformations"]["stop_words"]: 
                print("removing stop words")
                with open('stopwords.json', 'r') as file:
                    stopwords = loads(file.read())  # list of dutch stopwords
                total_doc_corpora[-1]["text"] = TextModifier.remove_stop_words(total_doc_corpora[-1]["text"], stopwords)

            if flag_dict["text_transformations"]["lemmatization"]:
                print("applying lemmatization")
                total_doc_corpora[-1]["text"] = TextModifier.apply_lemmatization(total_doc_corpora[-1]["text"])

            if flag_dict["text_transformations"]["stemming"]:
                print("applying stemming")
                total_doc_corpora[-1]["text"] = TextModifier.apply_stemming(total_doc_corpora[-1]["text"])
        return total_doc_corpora

    def _load(self, dest_str: str, data):
        """
        TODO: Docstring
        """
        json_str = jsonpickle.encode(data)
        with open (dest_str, "w") as f:
            f.write(json_str)

class AggregateETL(BaseETL):
    """ TODO: DOCSTRING """
    def __init__(self, data_dir: str, dest_dir: str):
        super().__init__(data_dir, dest_dir)

    def _transform(self, data: list[Document]):
        """ TODO: DOCSTRING """
        #reasignment for name clarity
        docx_files = data 
        number_of_sections=len(docx_files)
        #TODO: Check len for out of range
        master = data[0][0] #TODO: Might want to select master file manually using param
        composer = Composer(master)
        for i in range(1, number_of_sections):
            doc_temp = docx_files[i][0]
            composer.append(doc_temp)
        #composer.save("combined_file.docx")
        return composer

    def _load(self, dest_str: str, data: Composer):
        """ TODO: Docstring """
        data.save(dest_str)